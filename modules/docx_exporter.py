from pathlib import Path
import re

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

class DocxExporter:

    def __init__(self, resume_json):

        self.resume = resume_json
        self.document = Document()

        self.output_folder = Path("outputs")
        self.output_folder.mkdir(exist_ok=True)

        self.setup_document()

    # =====================================================
    # Document Styling
    # =====================================================

    def setup_document(self):

        section = self.document.sections[0]

        section.top_margin = Inches(0.4)
        section.bottom_margin = Inches(0.4)
        section.left_margin = Inches(0.4)
        section.right_margin = Inches(0.4)

        self.set_default_font()

    def set_default_font(self):

        normal = self.document.styles["Normal"]

        normal.font.name = "Calibri"
        normal.font.size = Pt(10.5)

        fmt = normal.paragraph_format
        fmt.space_before = Pt(0)
        fmt.space_after = Pt(2)
        fmt.line_spacing = 1

    # =====================================================
    # Utility Functions
    # =====================================================

    def add_horizontal_line(self):

        paragraph = self.document.add_paragraph()

        pPr = paragraph._p.get_or_add_pPr()

        pBdr = OxmlElement("w:pBdr")

        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")  # Line style
        bottom.set(qn("w:sz"), "4")      # Line thickness
        bottom.set(qn("w:space"), "0.5")  # Space between line and text
        bottom.set(qn("w:color"), "auto")  # Line color

        pBdr.append(bottom)
        pPr.append(pBdr)

    def add_section_heading(self, title):

        heading = self.document.add_paragraph()

        heading.paragraph_format.space_before = Pt(3)
        heading.paragraph_format.space_after = Pt(0)

        run = heading.add_run(title.upper())

        run.bold = True
        run.font.size = Pt(12)

        self.add_horizontal_line()
        # Divider line underneath

       # divider = self.document.add_paragraph()
        #divider.paragraph_format.space_before = Pt(0)
        #divider.paragraph_format.space_after = Pt(4)

        #divider.add_run("_" * 115)

    def add_spacer(self, size=4):

        paragraph = self.document.add_paragraph()

        paragraph.paragraph_format.space_after = Pt(size)
        
    # =====================================================
    # Header
    # =====================================================

    def add_header(self):

        personal = self.resume.get("personal_information", {})

        name = personal.get("name", "").strip()
        email = personal.get("email", "").strip()
        phone = personal.get("phone", "").strip()
        linkedin = personal.get("linkedin", "").strip()
        github = personal.get("github", "").strip()

        headline = self.resume.get("headline", "").strip()

        # ---------------------------------------
        # Candidate Name
        # ---------------------------------------

        heading = self.document.add_paragraph()

        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        heading.paragraph_format.space_after = Pt(2)

        run = heading.add_run(name.upper())

        run.bold = True
        run.font.size = Pt(22)

        # ---------------------------------------
        # Headline
        # ---------------------------------------

        if headline:

            p = self.document.add_paragraph()

            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            p.paragraph_format.space_after = Pt(4)

            run = p.add_run(headline)

            run.italic = True
            run.font.size = Pt(10)

        # ---------------------------------------
        # Contact Information
        # ---------------------------------------

        contact = []

        if email:
            contact.append(email)

        if phone:
            contact.append(phone)

        if linkedin:
            contact.append(linkedin)

        if github:
            contact.append(github)

        if contact:

            p = self.document.add_paragraph()

            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            p.paragraph_format.space_after = Pt(6)

            p.add_run(" | ".join(contact))

        # ---------------------------------------
        # Divider
        # ---------------------------------------

        self.add_horizontal_line()
        
    # =====================================================
    # Professional Summary
    # =====================================================

    def add_summary(self):

        summary = self.resume.get("summary", "").strip()

        if not summary:
            return

        self.add_section_heading("Profile")

        paragraph = self.document.add_paragraph()

        paragraph.paragraph_format.space_after = Pt(6)
        paragraph.paragraph_format.line_spacing = 1

        run = paragraph.add_run(summary)

        run.font.size = Pt(9)
        
    # =====================================================
    # Technical Skills
    # =====================================================

    def add_skills(self):

        skills = self.resume.get("skills", {})

        if not skills:
            return

        self.add_section_heading("Technical Skills")

        for category, items in skills.items():

            if not items:
                continue

            category_name = category.replace("_", " ").title()

            paragraph = self.document.add_paragraph(style="List Bullet")

            paragraph.paragraph_format.space_after = Pt(2)

            title = paragraph.add_run(f"{category_name}: ")

            title.bold = True
            title.font.size = Pt(10.5)

            skills_text = ", ".join(items)

            run = paragraph.add_run(skills_text)

            run.font.size = Pt(10.5)   
            
    # =====================================================
    # Professional Experience
    # =====================================================

    def add_experience(self):

        experience = self.resume.get("experience", [])

        if not experience:
            return

        self.add_section_heading("Professional Experience")

        for job in experience:

            title = job.get("job_title", "").strip()
            company = job.get("company", "").strip()
            location = job.get("location", "").strip()

            start = str(job.get("start_date") or "").strip()
            end = str(job.get("end_date") or "").strip()

            # --------------------------------------------
            # Job Heading
            # --------------------------------------------

            header = self.document.add_paragraph()

            header.paragraph_format.space_after = Pt(2)

            left = f"{title}"

            if company:
                left += f" | {company}"

            if location:
                left += f" | {location}"

            run = header.add_run(left)

            run.bold = True
            run.font.size = Pt(11)

            # Dates (right side)

            if start or end:

                header.add_run("\t")

                dates = header.add_run(f"{start} - {end}")

                dates.italic = True
                dates.font.size = Pt(10)

            # --------------------------------------------
            # Responsibilities
            # --------------------------------------------

            responsibilities = job.get("responsibilities", [])

            for responsibility in responsibilities[:4]:

                bullet = self.document.add_paragraph(
                    style="List Bullet"
                )

                bullet.paragraph_format.space_after = Pt(1)

                bullet.add_run(responsibility)

            self.add_spacer(3)
            
    # =====================================================
    # Projects
    # =====================================================

    def add_projects(self):

        projects = self.resume.get("projects", [])

        if not projects:
            return

        self.add_section_heading("Projects")

        # Show only the first 4 projects
        for project in projects[:4]:

            title = project.get("title", "").strip()
            technologies = project.get("technologies", [])
            description = project.get("description", "").strip()

            # --------------------------------------------
            # Project Title
            # --------------------------------------------

            header = self.document.add_paragraph()

            header.paragraph_format.space_after = Pt(2)

            run = header.add_run(title)

            run.bold = True
            run.font.size = Pt(11)

            if technologies:

                tech = ", ".join(technologies)

                header.add_run(" | ")

                tech_run = header.add_run(tech)

                tech_run.italic = True
                tech_run.font.size = Pt(10)

            # --------------------------------------------
            # Description
            # --------------------------------------------

            if description:

                bullet = self.document.add_paragraph(
                    style="List Bullet"
                )

                bullet.paragraph_format.space_after = Pt(1)

                bullet.add_run(description)

            self.add_spacer(3)
            
    # =====================================================
    # Education
    # =====================================================

    def add_education(self):

        education = self.resume.get("education", [])

        if not education:
            return

        self.add_section_heading("Education")

        for edu in education:

            degree = edu.get("degree", "").strip()
            institution = edu.get("institution", "").strip()
            dates = edu.get("dates", "").strip()
            description = edu.get("description", "").strip()

            # --------------------------------------------
            # Education Header
            # --------------------------------------------

            header = self.document.add_paragraph()

            header.paragraph_format.space_after = Pt(2)

            run = header.add_run(degree)

            run.bold = True
            run.font.size = Pt(11)

            if institution:

                header.add_run(" | ")

                uni = header.add_run(institution)

                uni.italic = True

            if dates:

                header.add_run(" | ")

                header.add_run(dates)

            # --------------------------------------------
            # Description
            # --------------------------------------------

            if description:

                bullet = self.document.add_paragraph(
                    style="List Bullet"
                )

                bullet.paragraph_format.space_after = Pt(1)

                bullet.add_run(description)

            self.add_spacer(3)
            
    # =====================================================
    # Languages
    # =====================================================

    def add_languages(self):

        languages = self.resume.get("languages", [])

        if not languages:
            return

        self.add_section_heading("Languages")

        paragraph = self.document.add_paragraph()

        paragraph.paragraph_format.space_after = Pt(4)

        first = True

        for language in languages:

            name = language.get("language", "").strip()
            proficiency = language.get("proficiency", "").strip()

            if not first:
                paragraph.add_run(" • ")

            text = name

            if proficiency:
                text += f" ({proficiency})"

            paragraph.add_run(text)

            first = False
            
    # =====================================================
    # Certifications
    # =====================================================

    def add_certifications(self):

        certifications = self.resume.get("certifications", [])

        if not certifications:
            return

        self.add_section_heading("Certifications")

        for cert in certifications[:4]:

            bullet = self.document.add_paragraph(
                style="List Bullet"
            )

            bullet.paragraph_format.space_after = Pt(1)

            bullet.add_run(cert)
            
    # =====================================================
    # Save Document
    # =====================================================

    def save_document(self):

        personal = self.resume.get("personal_information", {})

        name = personal.get("name", "Resume").strip()

        # Remove invalid filename characters

        safe_name = re.sub(r'[\\/*?:"<>|]', "", name)

        safe_name = safe_name.replace(" ", "_")

        filename = f"{safe_name}_Tailored_Resume.docx"

        output_file = self.output_folder / filename

        self.document.save(output_file)

        return output_file


    # =====================================================
    # Export Resume
    # =====================================================

    def export(self):

        self.add_header()

        self.add_summary()

        self.add_skills()

        self.add_experience()

        self.add_projects()

        self.add_education()

        self.add_certifications()

        self.add_languages()

        return self.save_document()